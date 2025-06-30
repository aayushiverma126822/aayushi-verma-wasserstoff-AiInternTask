import os
import docx
from langchain.schema import Document
from langchain.document_loaders import PyPDFLoader, TextLoader

def read_docx(file):
    doc = docx.Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    metadata = {"source": file.name}
    return [Document(page_content=text, metadata=metadata)]

def load_documents(uploaded_files):
    documents = []
    for file in uploaded_files:
        file_name = file.name
        if file_name.endswith(".pdf"):
            loader = PyPDFLoader(file)
            documents.extend(loader.load())
        elif file_name.endswith(".txt"):
            loader = TextLoader(file)
            documents.extend(loader.load())
        elif file_name.endswith(".docx"):
            documents.extend(read_docx(file))
        else:
            continue
    return documents
